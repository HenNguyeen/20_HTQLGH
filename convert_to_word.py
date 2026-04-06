#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown to Word Document (.docx)
Requires: python-docx library
Install: pip install python-docx
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Error: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)

def convert_markdown_to_docx(markdown_file, output_file):
    """Convert Markdown to Word Document"""
    
    # Check file exists
    if not os.path.exists(markdown_file):
        print(f"❌ File not found: {markdown_file}")
        return False
    
    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create Word document
    doc = Document()
    
    # Set margins (1 inch = 1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process markdown
    code_block = False
    code_lines = []
    table_buffer = []
    
    for line in lines:
        line = line.rstrip('\n\r')
        
        # Headers
        if line.startswith('# '):
            heading = line.lstrip('# ').strip()
            p = doc.add_paragraph(heading, style='Heading 1')
            p.runs[0].font.size = Pt(24)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0x36, 0x60, 0x92)
        
        elif line.startswith('## '):
            heading = line.lstrip('## ').strip()
            p = doc.add_paragraph(heading, style='Heading 2')
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        
        elif line.startswith('### '):
            heading = line.lstrip('### ').strip()
            p = doc.add_paragraph(heading, style='Heading 3')
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True
        
        # Code blocks
        elif line.startswith('```'):
            if code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text, style='Normal')
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                
                # Add background shading
                from docx.oxml import parse_xml
                shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
                p._element.get_or_add_pPr().append(shading_elm)
                
                code_block = False
                code_lines = []
            else:
                code_block = True
        
        elif code_block:
            code_lines.append(line)
        
        # Tables
        elif line.startswith('|') and '|' in line:
            cells = [c.strip() for c in line.split('|') if c.strip()]
            table_buffer.append(cells)
            
            # Separator line - create table after
            if line.replace('-', '').replace('|', '').replace(':', '').strip() == '':
                if table_buffer and len(table_buffer) > 1:
                    # Create table
                    table = doc.add_table(rows=len(table_buffer), cols=len(table_buffer[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row_data in enumerate(table_buffer):
                        row = table.rows[i]
                        for j, cell_data in enumerate(row_data):
                            cell = row.cells[j]
                            cell.text = cell_data
                            # Format header row
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                    table_buffer = []
        
        # Bullet points
        elif line.startswith('- '):
            text = line.lstrip('- ').strip()
            doc.add_paragraph(text, style='List Bullet')
        
        elif line.startswith('  - '):
            text = line.lstrip('  - ').strip()
            doc.add_paragraph(text, style='List Bullet 2')
        
        # Horizontal rule
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)
            from docx.oxml import parse_xml
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>')
            pPr.append(pBdr)
        
        # Regular text
        elif line.strip():
            # Remove markdown formatting
            text = line
            text = text.replace('**', '').replace('*', '')
            text = text.replace('`', '')
            text = text.replace('[', '').replace(']', '')
            
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(11)
        
        # Blank line
        else:
            doc.add_paragraph()
    
    # Save document
    doc.save(output_file)
    print(f"✅ Successfully exported to: {output_file}")
    print(f"📄 File size: {os.path.getsize(output_file) / 1024:.1f} KB")
    print(f"📍 Full path: {os.path.abspath(output_file)}")
    return True

if __name__ == '__main__':
    # Get parameters from command line or use defaults
    markdown_file = sys.argv[1] if len(sys.argv) > 1 else 'DesignPatternAnalysis.md'
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'DesignPatternAnalysis.docx'
    
    # Convert
    success = convert_markdown_to_docx(markdown_file, output_file)
    sys.exit(0 if success else 1)
